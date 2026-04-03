"""
ATS-Aware Resume Optimizer.

Parses PDF/DOCX resumes, analyzes ATS compatibility, matches keywords,
and generates optimized, ATS-friendly resumes for specific jobs.
"""

from __future__ import annotations

import json
import logging
import re
import sqlite3
import time
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from src.ai.brain import AIBrain

logger = logging.getLogger(__name__)

# Optional imports for file parsing
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logger.warning("PyPDF2 not installed — PDF parsing disabled")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed — DOCX parsing disabled")


# ---------------------------------------------------------------------------
# Data models
# ---------------------------------------------------------------------------

class ResumeFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


@dataclass
class ResumeData:
    """Parsed resume data."""
    raw_text: str
    sections: dict[str, str] = field(default_factory=dict)
    skills: list[str] = field(default_factory=list)
    experience_years: float | None = None
    education: list[str] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    file_format: ResumeFormat = ResumeFormat.TXT
    file_path: str = ""


@dataclass
class ATSScore:
    """ATS compatibility analysis result."""
    score: int  # 0-100
    issues: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    suggestions: list[str] = field(default_factory=list)
    has_tables: bool = False
    has_images: bool = False
    has_headers_footers: bool = False
    font_issues: list[str] = field(default_factory=list)
    section_detected: dict[str, bool] = field(default_factory=dict)


@dataclass
class KeywordMatch:
    """Keyword matching analysis."""
    missing: list[str] = field(default_factory=list)
    weak: list[str] = field(default_factory=list)
    strong: list[str] = field(default_factory=list)
    match_percentage: float = 0.0
    total_keywords: int = 0


@dataclass
class OptimizedResume:
    """Result of resume optimization for a specific job."""
    original_text: str
    optimized_text: str
    changes: list[str] = field(default_factory=list)
    keywords_added: list[str] = field(default_factory=list)
    sections_reordered: bool = False
    ats_score_before: int = 0
    ats_score_after: int = 0


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

STANDARD_SECTIONS = [
    "summary", "objective", "profile",
    "experience", "work experience", "employment", "professional experience",
    "education", "academic",
    "skills", "technical skills", "competencies",
    "certifications", "certificates", "licenses",
    "projects", "achievements", "awards",
    "publications", "languages", "interests", "references",
]

COMMON_SKILLS = [
    "python", "java", "javascript", "typescript", "react", "angular", "vue",
    "node.js", "django", "flask", "spring", "spring boot",
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "ci/cd",
    "git", "agile", "scrum", "jira",
    "machine learning", "deep learning", "nlp", "computer vision",
    "data analysis", "data science", "statistics",
    "html", "css", "rest api", "graphql", "microservices",
    "linux", "bash", "shell scripting",
    "excel", "power bi", "tableau", "powerpoint",
    "project management", "team leadership", "communication",
]

ATS_PROBLEM_PATTERNS = [
    (r'<table|<td|<tr', "Contains HTML tables — ATS may not parse correctly"),
    (r'\|.*\|.*\|', "Contains markdown tables — use simple bullet lists instead"),
    (r'[^\x00-\x7F]', "Contains non-ASCII characters that may confuse ATS"),
    (r'(?i)(\.png|\.jpg|\.gif|\.svg)', "Contains image references — ATS cannot read images"),
    (r'(?i)header|footer', "May contain headers/footers — ATS often ignores these"),
    (r'(?i)column|multi.column', "Multi-column layout detected — ATS prefers single column"),
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _extract_keywords_from_jd(jd_text: str) -> list[str]:
    """Extract important keywords from a job description."""
    text_lower = jd_text.lower()
    found = [skill for skill in COMMON_SKILLS if skill in text_lower]

    caps = re.findall(r'\b[A-Z][A-Za-z+#.]{2,}\b', jd_text)
    found.extend([c for c in caps if len(c) > 2 and c.lower() not in {f.lower() for f in found}])

    parens = re.findall(r'\(([^)]+)\)', jd_text)
    found.extend([p.strip() for p in parens if len(p.strip()) > 2])

    seen: set[str] = set()
    unique: list[str] = []
    for k in found:
        kl = k.lower()
        if kl not in seen:
            seen.add(kl)
            unique.append(k)
    return unique[:50]


def _extract_sections(text: str) -> dict[str, str]:
    """Extract resume sections by looking for section headers."""
    sections: dict[str, str] = {}
    lines = text.split("\n")
    current_section = "header"
    current_content: list[str] = []

    for line in lines:
        stripped = line.strip().lower().rstrip(":")
        is_section = False
        for section_name in STANDARD_SECTIONS:
            if stripped == section_name or stripped.replace(" ", "") == section_name.replace(" ", ""):
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = section_name
                current_content = []
                is_section = True
                break
        if not is_section:
            current_content.append(line)

    if current_content:
        sections[current_section] = "\n".join(current_content).strip()
    return sections


def _extract_skills(text: str) -> list[str]:
    """Extract skills from resume text."""
    text_lower = text.lower()
    return list({skill for skill in COMMON_SKILLS if skill in text_lower})


def _extract_experience_years(text: str) -> float | None:
    """Try to extract years of experience from resume text."""
    patterns = [
        r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
        r'experience[:\s]*(\d+)\+?\s*years?',
        r'(\d+)\+?\s*yrs',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                return float(match.group(1))
            except ValueError:
                pass
    return None


# ---------------------------------------------------------------------------
# ResumeOptimizer
# ---------------------------------------------------------------------------

class ResumeOptimizer:
    """ATS-aware resume optimizer for job applications.

    Parses resumes (PDF/DOCX/TXT), analyzes ATS compatibility, matches
    keywords against job descriptions, and generates optimized resumes.

    Optionally uses an AI brain for enhanced optimization.
    """

    def __init__(self, db_path: str | None = None, brain: "AIBrain | None" = None):
        self._db_path = db_path or str(
            Path.home() / ".openclaw/workspace/job-auto-apply/data/resumes.db"
        )
        Path(self._db_path).parent.mkdir(parents=True, exist_ok=True)
        self._brain = brain
        self._init_db()

    def _init_db(self) -> None:
        with sqlite3.connect(self._db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS resume_versions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    application_id TEXT,
                    job_id TEXT,
                    original_path TEXT,
                    optimized_text TEXT,
                    template TEXT,
                    keywords_added TEXT,
                    created_at REAL
                )
            """)

    # --- Parsing -----------------------------------------------------------

    def parse_resume(self, file_path: str) -> ResumeData:
        """Parse a resume file (PDF, DOCX, or TXT)."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            raw_text = self._parse_pdf(str(path))
            fmt = ResumeFormat.PDF
        elif suffix in (".docx", ".doc"):
            raw_text = self._parse_docx(str(path))
            fmt = ResumeFormat.DOCX
        elif suffix == ".txt":
            raw_text = path.read_text(encoding="utf-8", errors="replace")
            fmt = ResumeFormat.TXT
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        sections = _extract_sections(raw_text)
        skills = _extract_skills(raw_text)
        exp_years = _extract_experience_years(raw_text)

        return ResumeData(
            raw_text=raw_text,
            sections=sections,
            skills=skills,
            experience_years=exp_years,
            education=self._extract_education(raw_text),
            certifications=self._extract_certifications(raw_text),
            file_format=fmt,
            file_path=str(path),
        )

    def _parse_pdf(self, file_path: str) -> str:
        if not HAS_PDF:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)

    def _parse_docx(self, file_path: str) -> str:
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    def _extract_education(self, text: str) -> list[str]:
        degree_patterns = [
            r"(?i)(b\.?tech|b\.?e\.?|bachelor|b\.?sc|b\.?a\.?|b\.?com|m\.?tech|m\.?e\.?|"
            r"master|m\.?sc|m\.?a\.?|m\.?com|mba|phd|doctorate|diploma)[^.]*\.",
        ]
        education = []
        for pattern in degree_patterns:
            matches = re.findall(pattern, text)
            education.extend(matches)
        return education

    def _extract_certifications(self, text: str) -> list[str]:
        cert_patterns = [
            r"(?i)(aws|azure|gcp|google|microsoft|oracle|cisco|comptia|pmp|scrum)[^.]*certif\w+",
            r"(?i)certif\w+\s*:\s*([^\n]+)",
        ]
        certs = []
        for pattern in cert_patterns:
            matches = re.findall(pattern, text)
            certs.extend([m if isinstance(m, str) else m[0] for m in matches])
        return certs

    # --- ATS analysis ------------------------------------------------------

    def analyze_ats_compatibility(self, resume_text: str) -> ATSScore:
        """Analyze resume text for ATS compatibility."""
        score = 100
        issues: list[str] = []
        warnings: list[str] = []
        suggestions: list[str] = []
        has_tables = False
        has_images = False

        for pattern, message in ATS_PROBLEM_PATTERNS:
            if re.search(pattern, resume_text):
                if "table" in message.lower():
                    has_tables = True
                if "image" in message.lower():
                    has_images = True
                issues.append(message)
                score -= 10

        sections_found = _extract_sections(resume_text)
        section_detected: dict[str, bool] = {}
        for s in ("experience", "education", "skills"):
            detected = s in sections_found or s.replace(" ", "") in "".join(sections_found.keys()).lower().replace(" ", "")
            section_detected[s] = detected
            if not detected:
                warnings.append(f"Section '{s}' not clearly detected by ATS parser")
                score -= 5

        word_count = len(resume_text.split())
        if word_count < 200:
            warnings.append("Resume seems very short (<200 words)")
            score -= 15
        elif word_count > 2000:
            warnings.append("Resume is long (>2000 words) — consider trimming for ATS")
            score -= 5

        if not re.search(r'[\w.+-]+@[\w-]+\.[\w.]+', resume_text):
            warnings.append("No email address detected")
            score -= 10
        if not re.search(r'(\+?\d{1,3}[-.\s]?)?\d{10}', resume_text):
            warnings.append("No phone number detected")
            score -= 5

        if not sections_found:
            suggestions.append("Add clear section headers (Experience, Education, Skills)")
        if word_count > 1500:
            suggestions.append("Keep resume to 1-2 pages for optimal ATS parsing")

        return ATSScore(
            score=max(0, min(100, score)),
            issues=issues,
            warnings=warnings,
            suggestions=suggestions,
            has_tables=has_tables,
            has_images=has_images,
            section_detected=section_detected,
        )

    # --- Keyword matching --------------------------------------------------

    def match_keywords(self, jd_text: str, resume_text: str) -> KeywordMatch:
        """Match keywords between job description and resume."""
        jd_keywords = _extract_keywords_from_jd(jd_text)
        resume_lower = resume_text.lower()

        if not jd_keywords:
            return KeywordMatch(total_keywords=0, match_percentage=100.0)

        strong: list[str] = []
        weak: list[str] = []
        missing: list[str] = []

        for kw in jd_keywords:
            kw_lower = kw.lower()
            if kw_lower in resume_lower:
                if resume_lower.count(kw_lower) >= 2:
                    strong.append(kw)
                else:
                    weak.append(kw)
            else:
                missing.append(kw)

        total = len(jd_keywords)
        match_pct = (len(strong) + len(weak) * 0.5) / total * 100 if total > 0 else 0

        return KeywordMatch(
            missing=missing,
            weak=weak,
            strong=strong,
            match_percentage=round(match_pct, 1),
            total_keywords=total,
        )

    # --- Optimization ------------------------------------------------------

    def optimize_for_job(self, jd_text: str, resume_data: ResumeData) -> OptimizedResume:
        """Optimize a resume for a specific job description.

        If an AI brain is available, uses it for enhanced optimization.
        Otherwise falls back to local keyword-based optimization.
        """
        if self._brain:
            return self._ai_optimize(jd_text, resume_data)
        return self._local_optimize(jd_text, resume_data)

    def _ai_optimize(self, jd_text: str, resume_data: ResumeData) -> OptimizedResume:
        """AI-enhanced resume optimization via brain."""
        original = resume_data.raw_text
        ats_before = self.analyze_ats_compatibility(original)

        prompt = (
            "Optimize this resume for the given job description. Return JSON with:\n"
            "- optimized_text: the full optimized resume text\n"
            "- changes: list of changes made\n"
            "- keywords_added: list of keywords added\n"
            "- sections_reordered: bool\n\n"
            f"JOB DESCRIPTION:\n{jd_text[:4000]}\n\n"
            f"RESUME:\n{original[:4000]}\n\n"
            "Respond with ONLY valid JSON."
        )

        response = self._brain._call_api([
            {"role": "system", "content": "You optimize resumes for ATS compatibility. Respond only with JSON."},
            {"role": "user", "content": prompt},
        ])

        if response:
            try:
                text = response.strip()
                if text.startswith("```"):
                    text = text.split("\n", 1)[1].rsplit("```", 1)[0].strip()
                result = json.loads(text)
                optimized_text = result.get("optimized_text", original)
                ats_after = self.analyze_ats_compatibility(optimized_text)
                return OptimizedResume(
                    original_text=original,
                    optimized_text=optimized_text,
                    changes=result.get("changes", []),
                    keywords_added=result.get("keywords_added", []),
                    sections_reordered=result.get("sections_reordered", False),
                    ats_score_before=ats_before.score,
                    ats_score_after=ats_after.score,
                )
            except (json.JSONDecodeError, KeyError, ValueError) as e:
                logger.warning(f"Failed to parse AI optimization response: {e}")

        # Fallback to local
        return self._local_optimize(jd_text, resume_data)

    def _local_optimize(self, jd_text: str, resume_data: ResumeData) -> OptimizedResume:
        """Local keyword-based resume optimization."""
        original = resume_data.raw_text
        changes: list[str] = []
        keywords_added: list[str] = []
        sections_reordered = False

        keyword_match = self.match_keywords(jd_text, original)
        ats_before = self.analyze_ats_compatibility(original)

        optimized_lines = original.split("\n")

        # 1. Add missing keywords
        if keyword_match.missing:
            skills_section_idx = None
            for i, line in enumerate(optimized_lines):
                if any(s in line.lower() for s in ("skills", "technical skills", "competencies")):
                    skills_section_idx = i
                    break

            relevant_missing = [
                kw for kw in keyword_match.missing[:10]
                if len(kw) > 2 and not kw[0].isupper()
            ]

            if relevant_missing and skills_section_idx is not None:
                optimized_lines.insert(skills_section_idx + 1, f"Additional: {', '.join(relevant_missing)}")
                keywords_added = relevant_missing
                changes.append(f"Added {len(relevant_missing)} missing keywords to skills section")
            elif relevant_missing:
                optimized_lines.extend(["", "SKILLS", ", ".join(relevant_missing)])
                keywords_added = relevant_missing
                changes.append(f"Added skills section with {len(relevant_missing)} keywords")

        # 2. Flag weak keywords
        if keyword_match.weak:
            changes.append(f"Identified {len(keyword_match.weak)} weak keywords to emphasize")

        # 3. Section reordering suggestions
        jd_lower = jd_text.lower()
        if "certification" in jd_lower or "certified" in jd_lower:
            sections = _extract_sections("\n".join(optimized_lines))
            if "certifications" in sections and list(sections.keys()).index("certifications") > 3:
                changes.append("Recommended: move Certifications section higher (JD emphasizes certifications)")
                sections_reordered = True

        if "education" in jd_lower and ("degree" in jd_lower or "bachelor" in jd_lower or "master" in jd_lower):
            changes.append("Recommended: highlight Education section (JD emphasizes degree requirements)")

        # 4. Clean ATS issues
        cleaned = "\n".join(optimized_lines)
        cleaned = re.sub(r'[^\x00-\x7F]+', ' ', cleaned)
        cleaned = re.sub(r'\t+', '    ', cleaned)
        cleaned = re.sub(r' {3,}', '  ', cleaned)
        if cleaned != original:
            changes.append("Cleaned non-ASCII characters and normalized whitespace")

        ats_after = self.analyze_ats_compatibility(cleaned)

        return OptimizedResume(
            original_text=original,
            optimized_text=cleaned,
            changes=changes,
            keywords_added=keywords_added,
            sections_reordered=sections_reordered,
            ats_score_before=ats_before.score,
            ats_score_after=ats_after.score,
        )

    # --- Formatting --------------------------------------------------------

    def generate_formatted_resume(
        self, optimized: OptimizedResume, template: str = "ats_optimized"
    ) -> bytes:
        """Generate a formatted resume as bytes."""
        text = optimized.optimized_text
        if template == "govt_format":
            text = "=" * 60 + "\nBIO-DATA / CURRICULUM VITAE\n" + "=" * 60 + "\n\n" + text
        elif template == "private_format":
            text = "-" * 40 + "\nRESUME\n" + "-" * 40 + "\n\n" + text
        else:
            text = re.sub(r'\n{3,}', '\n\n', text)
        return text.encode("utf-8")

    # --- Tracking ----------------------------------------------------------

    def track_resume_version(
        self,
        application_id: str,
        job_id: str,
        original_path: str,
        optimized_text: str,
        template: str,
        keywords_added: list[str],
    ) -> None:
        """Track which resume version was used for which application."""
        with sqlite3.connect(self._db_path) as conn:
            conn.execute(
                "INSERT INTO resume_versions "
                "(application_id, job_id, original_path, optimized_text, template, keywords_added, created_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?)",
                (application_id, job_id, original_path, optimized_text, template,
                 json.dumps(keywords_added), time.time()),
            )

    def get_resume_versions(self, job_id: str | None = None) -> list[dict[str, Any]]:
        """Get tracked resume versions, optionally filtered by job_id."""
        with sqlite3.connect(self._db_path) as conn:
            if job_id:
                rows = conn.execute(
                    "SELECT * FROM resume_versions WHERE job_id = ? ORDER BY created_at DESC",
                    (job_id,),
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT * FROM resume_versions ORDER BY created_at DESC LIMIT 50"
                ).fetchall()
            columns = [d[0] for d in conn.execute("SELECT * FROM resume_versions LIMIT 0").description]
            return [dict(zip(columns, row)) for row in rows]
